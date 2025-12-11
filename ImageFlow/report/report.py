import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Cm

import io

# 리포트 생성 함수
def CreateReport(model_train_states, save_path):
        
    plt.figure(figsize=(14,6))
    
    
    # Loss Graph 생성 
    plt.subplot(1, 2, 1)
    for state in model_train_states:
        
        # X축 : 학습 횟수
        epochs = range(1, len(state['train_history'])+1)
        
        # Y축 : 학습 및 검증 손실
        train_loss = [i[0] for i in state['train_history']]
        valid_loss = [i[0] for i in state['valid_history']]
        
        # 학습 및 검증 색상 통일
        lc = plt.plot(epochs, train_loss, linestyle='--', marker='.', label=f'{state["type"]} (Train)')
        plt.plot(epochs, valid_loss, linestyle='-' , marker='o', label=f'{state["type"]} (Valid)', color=lc[0].get_color())
    
    plt.title('Train & Validation Loss Graph')
    plt.xlabel('Epoch')
    plt.ylabel('Loss')
    
    plt.legend()
    plt.grid(True)
    
    
    # Accuracy Graph 생성
    plt.subplot(1, 2, 2)
    for state in model_train_states:
        
        # X축 : 학습 횟수
        epochs = range(1, len(state['train_history'])+1)
        
        # Y축 : 학습 및 검증 정확도
        train_acc = [i[1] for i in state['train_history']]
        valid_acc = [i[1] for i in state['valid_history']]
        
        # 학습 및 검증 색상 통일
        lc = plt.plot(epochs, train_acc, linestyle='--', marker='.', label=f'{state["type"]} (Train)')
        plt.plot(epochs, valid_acc, linestyle='-' , marker='o', label=f'{state["type"]} (Valid)', color=lc[0].get_color())
    
    plt.title('Train & Validation Accuracy Graph')
    plt.xlabel('Epoch')
    plt.ylabel('Accuracy')
    
    plt.legend()
    plt.grid(True)
    
    # 임시 이미지 저장 
    plt.tight_layout()
    buffer = io.BytesIO()
    
    plt.savefig(buffer, format='png', dpi=300)
    plt.close()
    
    # 검증 손실이 작은 모델을 최고 모델로 선정
    best_model = min(model_train_states, key=lambda x: x['valid_loss'])
    
    # 워드 페이지 구성 
    doc = Document()
    
    doc.add_heading('ImageFlow Training Report', 0)
    doc.add_heading('Best Model Information', 1)
    doc.add_paragraph(f'Based on the lowest validation loss, best model is: {best_model["save_name"]}')
    doc.add_paragraph(f'- Validation Loss: {best_model["valid_loss"]:.4f}')
    doc.add_paragraph(f'- Accuracy Loss: {best_model["valid_acc"]:.4f}')
    doc.add_paragraph(f'- Model checkpoint path: {best_model["save_dir"]}/{best_model["save_name"]}.onnx')
    doc.add_paragraph(f'- training class_to_idx: {best_model["class_to_idx"]}')
    
    doc.add_heading('Hyperparameters used:', 2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    header = table.rows[0].cells
    header[0].text = 'Parameter'
    header[1].text = 'Value'
    
    for k,v in best_model['params'].items():
        
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)
    
    doc.add_heading('Training Visualizations', 1)
    doc.add_picture(buffer, width=Cm(24))
    
    # 워드 저장
    doc.save(save_path)
    
    
    
    
    
        
        
        